from docx import Document
from docx.shared import Inches
import io


def build_docx(event,report_text,bod):

    doc=Document()

    doc.add_picture("logo.png",width=Inches(1.5))

    doc.add_heading("Rotaract Club of Bombay Airport",level=0)

    doc.add_heading(event["title"],level=1)

    doc.add_paragraph(f"Venue: {event['venue']}")
    doc.add_paragraph(f"Start Time: {event['start_time']}")
    doc.add_paragraph(f"End Time: {event['end_time']}")
    doc.add_paragraph(f"Chief Guest: {event['chief_guest']}")

    doc.add_heading("Event Report",level=2)

    doc.add_paragraph(report_text)

    doc.add_heading("Project Details",level=2)

    for k,v in bod.items():
        doc.add_paragraph(f"{k}: {v}")

    buffer=io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer