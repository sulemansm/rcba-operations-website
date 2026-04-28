"""
Rotaract Club of Bombay Airport — Event Reporting System
Roles: admin (super), secretariat (full admin), editor (view all + submit), director (own + submit)
Storage: Supabase (primary) with JSON fallback
Deployment: Render.com
"""

import base64
import io
import json
import os
import secrets
import smtplib
import urllib.parse
from datetime import date, datetime, timedelta
from email.message import EmailMessage

import requests as req
import streamlit as st
from dotenv import load_dotenv

from secrets_manager import (
    get_secret,
    get_oauth_redirect_uri,
    load_google_credentials,
    get_whitelisted_emails,
    has_google_credentials,
)
from auth import get_role, role_display_label, role_badge_class, can_submit_report
from report_handler import (
    save_report, load_reports, update_report, update_report_by_id,
    is_late, is_late_report, get_status,
)
from dashboard import (
    page_dashboard_secretariat,
    page_dashboard_editor,
    page_dashboard_director,
    page_admin,
    page_manage_members,
)

load_dotenv()

# ── secrets ───────────────────────────────────────────────────────────────────
GROQ_API_KEY    = get_secret("GROQ_API_KEY", "")
SENDER_EMAIL    = get_secret("SENDER_EMAIL", "")
SENDER_PASSWORD = get_secret("SENDER_PASSWORD", "")
SECRETARY_EMAIL = get_secret("SECRETARY_EMAIL", "")
REDIRECT_URI    = get_oauth_redirect_uri()

WHITELISTED_EMAILS = get_whitelisted_emails()

GOOGLE_AUTH_URL  = "https://accounts.google.com/o/oauth2/v2/auth"
GOOGLE_TOKEN_URL = "https://oauth2.googleapis.com/token"
GOOGLE_INFO_URL  = "https://www.googleapis.com/oauth2/v3/userinfo"

REPORT_SECTIONS       = ["Aim", "Execution", "Impact Analysis", "Follow Up and Feedback"]
EXECUTION_SUBSECTIONS = ["Pre-Event Work", "On-Day Work", "Post-Event Work"]

WORD_LIMITS = {
    "description": 80, "pre_event": 100,
    "on_day": 100, "post_event": 80, "outcome": 80,
}

RI_YEAR = "2024–2025"

# ── Google OAuth credentials ───────────────────────────────────────────────────
_gc                  = load_google_credentials()
GOOGLE_CLIENT_ID     = _gc.get("client_id", "")
GOOGLE_CLIENT_SECRET = _gc.get("client_secret", "")


# ═══════════════════════════════════════════════════════════════════════════════
# LOGO
# ═══════════════════════════════════════════════════════════════════════════════

def _logo_b64() -> str | None:
    app_dir = os.path.dirname(os.path.abspath(__file__))
    for nm in ["logo", "club_logo", "rcba_logo", "rotaract_logo"]:
        for ex in ["png", "jpg", "jpeg", "webp"]:
            p = os.path.join(app_dir, f"{nm}.{ex}")
            if os.path.exists(p):
                mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(ex, "image/png")
                return f"data:{mime};base64," + base64.b64encode(open(p, "rb").read()).decode()
    return None


def _logo_path() -> str | None:
    app_dir = os.path.dirname(os.path.abspath(__file__))
    for nm in ["logo", "club_logo", "rcba_logo", "rotaract_logo"]:
        for ex in ["png", "jpg", "jpeg"]:
            p = os.path.join(app_dir, f"{nm}.{ex}")
            if os.path.exists(p):
                return p
    return None


LOGO_SRC = _logo_b64()


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ═══════════════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="RCBA Event Reporter",
    page_icon="🌿",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;500;600;700;800&family=Outfit:wght@300;400;500;600&display=swap');

:root {
    --teal:      #00C9B1;
    --teal-glow: rgba(0,201,177,0.18);
    --teal-dark: #00A896;
    --gold:      #FFB347;
    --gold-glow: rgba(255,179,71,0.15);
    --navy:      #0A0F1E;
    --surface:   #111827;
    --surface2:  #1A2235;
    --surface3:  #1F2D42;
    --border:    rgba(255,255,255,0.08);
    --border-act:rgba(0,201,177,0.4);
    --text:      #E8EDF5;
    --muted:     #7A8BA0;
    --danger:    #FF6B6B;
    --success:   #4ECDC4;
    --admin:     #A78BFA;
    --admin-glow:rgba(167,139,250,0.15);
}

html, body, .stApp { background: var(--navy) !important; font-family: 'Outfit', sans-serif; color: var(--text); }
.stApp::before {
    content:''; position:fixed; inset:0;
    background: radial-gradient(ellipse 80% 50% at 20% 20%,rgba(0,201,177,0.06) 0%,transparent 60%),
                radial-gradient(ellipse 60% 40% at 80% 80%,rgba(255,179,71,0.05) 0%,transparent 60%);
    pointer-events:none; z-index:0;
}

/* HEADER */
.rcba-header { position:relative; border-radius:20px; overflow:hidden; margin-bottom:0; background:linear-gradient(135deg,#111827 0%,#1A2235 50%,#111827 100%); border:1px solid var(--border); box-shadow:0 0 0 1px rgba(0,201,177,0.1),0 20px 60px rgba(0,0,0,0.5); }
.rcba-header::before { content:''; position:absolute; inset:0; background:linear-gradient(135deg,rgba(0,201,177,0.05) 0%,transparent 50%,rgba(255,179,71,0.03) 100%); pointer-events:none; }
.rcba-header-stripe { height:3px; background:linear-gradient(90deg,var(--teal),var(--gold),var(--teal)); background-size:200% 100%; animation:shimmer 3s linear infinite; }
@keyframes shimmer { 0%{background-position:200% 0} 100%{background-position:-200% 0} }
.rcba-header-body { display:flex; align-items:center; gap:1.4rem; padding:1.4rem 1.8rem 1rem; position:relative; }
.rcba-logo-wrap { position:relative; flex-shrink:0; width:66px; height:66px; }
.rcba-logo-wrap img { width:66px; height:66px; object-fit:contain; display:block; filter:drop-shadow(0 0 14px rgba(0,201,177,0.35)); }
.rcba-logo-ring  { position:absolute; inset:-5px;  border-radius:50%; border:1.5px solid rgba(0,201,177,0.3);  animation:pulse-ring 2.5s ease-in-out infinite; }
.rcba-logo-ring2 { position:absolute; inset:-10px; border-radius:50%; border:1px   solid rgba(255,179,71,0.15); animation:pulse-ring 2.5s ease-in-out infinite 0.5s; }
@keyframes pulse-ring { 0%,100%{opacity:0.4;transform:scale(1)} 50%{opacity:0.9;transform:scale(1.03)} }
.rcba-header-text h1 { font-family:'Syne',sans-serif; color:#fff; font-size:1.35rem; margin:0 0 0.1rem; letter-spacing:-0.01em; font-weight:700; }
.rcba-header-text .sub { color:var(--muted); font-size:0.68rem; letter-spacing:0.12em; text-transform:uppercase; margin:0; }
.rcba-header-text .tagline { display:inline-flex; align-items:center; gap:0.4rem; color:var(--teal); font-size:0.68rem; letter-spacing:0.1em; text-transform:uppercase; margin:0.3rem 0 0; font-weight:600; font-family:'Syne',sans-serif; }
.rcba-header-text .tagline::before { content:''; width:14px; height:1.5px; background:var(--teal); display:inline-block; }

/* NAVBAR */
.rcba-nav { background:var(--surface); border:1px solid var(--border); border-top:none; border-radius:0 0 16px 16px; padding:0 1.8rem; display:flex; align-items:center; justify-content:space-between; margin-bottom:1.8rem; box-shadow:0 8px 32px rgba(0,0,0,0.3); }
.rcba-nav-links { display:flex; align-items:center; gap:0.2rem; }
.rcba-nav-btn { font-family:'Syne',sans-serif; font-size:0.78rem; font-weight:600; letter-spacing:0.06em; text-transform:uppercase; color:var(--muted); padding:0.9rem 1.1rem; border:none; background:transparent; cursor:pointer; transition:all .2s; border-bottom:2px solid transparent; white-space:nowrap; }
.rcba-nav-btn:hover { color:var(--text); }
.rcba-nav-btn.active { color:var(--teal); border-bottom-color:var(--teal); }
.rcba-nav-right { display:flex; align-items:center; gap:0.8rem; }

/* ROLE BADGES */
.rcba-role-badge { font-family:'Syne',sans-serif; font-size:0.62rem; font-weight:700; letter-spacing:0.1em; text-transform:uppercase; padding:0.2rem 0.6rem; border-radius:20px; }
.rcba-role-badge.member      { background:rgba(0,201,177,0.12);  color:var(--teal);  border:1px solid rgba(0,201,177,0.25); }
.rcba-role-badge.admin       { background:rgba(167,139,250,0.12);color:var(--admin); border:1px solid rgba(167,139,250,0.3); }
.rcba-role-badge.secretariat { background:rgba(167,139,250,0.12);color:var(--admin); border:1px solid rgba(167,139,250,0.3); }
.rcba-role-badge.editor      { background:rgba(255,179,71,0.12); color:var(--gold);  border:1px solid rgba(255,179,71,0.3); }
.rcba-role-badge.director    { background:rgba(0,201,177,0.12);  color:var(--teal);  border:1px solid rgba(0,201,177,0.25); }

.rcba-nav-avatar      { width:30px; height:30px; border-radius:50%; border:2px solid var(--teal); object-fit:cover; cursor:pointer; }
.rcba-nav-avatar-init { width:30px; height:30px; border-radius:50%; background:linear-gradient(135deg,var(--teal),var(--teal-dark)); display:inline-flex; align-items:center; justify-content:center; color:var(--navy); font-size:0.75rem; font-weight:800; flex-shrink:0; font-family:'Syne',sans-serif; }

/* MISC */
.rcba-step { display:flex; align-items:center; gap:0.8rem; margin:2rem 0 0.9rem; }
.rcba-step-pill { background:linear-gradient(135deg,var(--teal),var(--teal-dark)); color:var(--navy); font-size:0.64rem; font-weight:700; letter-spacing:0.14em; text-transform:uppercase; padding:0.25rem 0.75rem; border-radius:20px; white-space:nowrap; font-family:'Syne',sans-serif; box-shadow:0 0 12px var(--teal-glow); }
.rcba-step-title { font-family:'Syne',sans-serif; font-size:1.05rem; color:var(--text); font-weight:600; margin:0; letter-spacing:-0.01em; }
.rcba-page-title { font-family:'Syne',sans-serif; font-size:1.4rem; font-weight:700; color:var(--text); margin:0 0 0.3rem; }
.rcba-page-sub   { font-size:0.82rem; color:var(--muted); margin:0 0 1.5rem; }
.rcba-card { background:var(--surface); border:1px solid var(--border); border-radius:16px; padding:1.5rem 1.8rem; margin-bottom:0.8rem; box-shadow:0 4px 24px rgba(0,0,0,0.3); position:relative; overflow:hidden; }
.rcba-card::before { content:''; position:absolute; top:0; left:0; right:0; height:1px; background:linear-gradient(90deg,transparent,rgba(0,201,177,0.3),transparent); }
.rcba-hr { height:1px; background:linear-gradient(90deg,transparent,rgba(255,255,255,0.06) 30%,rgba(255,255,255,0.06) 70%,transparent); margin:1.8rem 0; border:none; }
.word-counter { font-size:0.7rem; color:var(--muted); text-align:right; margin-top:0.2rem; }
.word-counter.over { color:var(--danger); font-weight:600; }
.word-counter.near { color:var(--gold); }
.rcba-section-label { font-family:'Syne',sans-serif; font-size:0.68rem; font-weight:700; letter-spacing:0.14em; text-transform:uppercase; color:var(--teal); margin:0 0 0.3rem; display:flex; align-items:center; gap:0.5rem; }
.rcba-section-label::before { content:''; width:18px; height:2px; background:var(--teal); display:inline-block; border-radius:2px; }
.rcba-auto-field { background:var(--surface3); border:1px dashed rgba(0,201,177,0.25); border-radius:10px; padding:0.6rem 1rem; margin-top:0.3rem; display:flex; flex-direction:column; }
.rcba-auto-field .label { font-size:0.68rem; color:var(--muted); font-weight:500; margin-bottom:0.2rem; text-transform:uppercase; letter-spacing:0.06em; }
.rcba-auto-field .val   { font-family:'Syne',sans-serif; font-size:1rem; color:var(--teal); font-weight:700; }
.rcba-auto-negative .val { color:var(--danger) !important; }
.rcba-edit-banner { background:linear-gradient(135deg,rgba(0,201,177,0.08),rgba(0,201,177,0.02)); border:1px solid rgba(0,201,177,0.15); border-radius:12px; padding:1rem 1.2rem; margin-bottom:1rem; }
.rcba-edit-banner h4 { font-family:'Syne',sans-serif; color:var(--text); font-size:0.88rem; font-weight:700; margin:0 0 0.2rem; }
.rcba-edit-banner p  { font-size:0.76rem; color:var(--muted); margin:0; line-height:1.5; }
.rcba-bod { background:linear-gradient(135deg,rgba(255,179,71,0.06),rgba(255,179,71,0.02)); border:1px solid rgba(255,179,71,0.2); border-left:3px solid var(--gold); border-radius:0 14px 14px 0; padding:1.1rem 1.4rem 0.5rem; margin-bottom:1.2rem; }
.rcba-bod-title { font-family:'Syne',sans-serif; font-size:0.92rem; color:var(--text); font-weight:700; margin-bottom:0.2rem; }
.rcba-bod-sub   { font-size:0.76rem; color:var(--muted); margin-bottom:0; line-height:1.5; }

/* PROFILE */
.rcba-profile-hero { background:linear-gradient(135deg,var(--surface2),var(--surface3)); border:1px solid var(--border); border-radius:20px; padding:2.5rem 2rem; text-align:center; margin-bottom:1.2rem; position:relative; overflow:hidden; }
.rcba-profile-hero::before { content:''; position:absolute; top:0; left:0; right:0; height:2px; background:linear-gradient(90deg,var(--teal),var(--gold),var(--teal)); background-size:200% 100%; animation:shimmer 3s linear infinite; }
.rcba-profile-avatar-lg { width:80px; height:80px; border-radius:50%; margin:0 auto 1rem; border:3px solid var(--teal); object-fit:cover; box-shadow:0 0 30px var(--teal-glow); }
.rcba-profile-avatar-lg-init { width:80px; height:80px; border-radius:50%; margin:0 auto 1rem; background:linear-gradient(135deg,var(--teal),var(--teal-dark)); display:flex; align-items:center; justify-content:center; color:var(--navy); font-size:2rem; font-weight:800; font-family:'Syne',sans-serif; box-shadow:0 0 30px var(--teal-glow); }
.rcba-profile-name  { font-family:'Syne',sans-serif; font-size:1.4rem; font-weight:700; color:var(--text); margin:0 0 0.25rem; }
.rcba-profile-email { font-size:0.85rem; color:var(--muted); margin:0 0 1rem; }
.rcba-profile-stat  { display:inline-flex; flex-direction:column; align-items:center; gap:0.15rem; padding:0 1.5rem; }
.rcba-profile-stat .n { font-family:'Syne',sans-serif; font-size:1.5rem; font-weight:700; color:var(--teal); }
.rcba-profile-stat .l { font-size:0.7rem; color:var(--muted); text-transform:uppercase; letter-spacing:0.08em; }
.rcba-profile-stats { display:flex; justify-content:center; gap:0; border-top:1px solid var(--border); padding-top:1.2rem; margin-top:1rem; }
.rcba-profile-stats .divider { width:1px; background:var(--border); margin:0; }
.rcba-info-row { display:flex; align-items:center; justify-content:space-between; padding:0.85rem 1rem; border-bottom:1px solid rgba(255,255,255,0.04); }
.rcba-info-row:last-child { border-bottom:none; }
.rcba-info-row .key { font-size:0.78rem; color:var(--muted); font-weight:500; text-transform:uppercase; letter-spacing:0.06em; }
.rcba-info-row .val { font-size:0.88rem; color:var(--text); font-weight:500; }

/* LOGIN */
.rcba-login-outer { max-width:420px; margin:3rem auto 0; }
.rcba-login-card { background:var(--surface); border:1px solid var(--border); border-radius:20px; padding:2.5rem 2rem 2.2rem; text-align:center; box-shadow:0 0 0 1px rgba(0,201,177,0.08),0 24px 80px rgba(0,0,0,0.6); position:relative; overflow:hidden; }
.rcba-login-card::before { content:''; position:absolute; top:0; left:0; right:0; height:2px; background:linear-gradient(90deg,var(--teal),var(--gold),var(--teal)); background-size:200% 100%; animation:shimmer 3s linear infinite; }
.rcba-login-card img { width:88px; height:88px; object-fit:contain; margin-bottom:1.2rem; filter:drop-shadow(0 0 20px rgba(0,201,177,0.3)); }
.rcba-login-card h2 { font-family:'Syne',sans-serif; color:var(--text); font-size:1.45rem; margin:0 0 0.3rem; font-weight:700; }
.rcba-login-card .club-tag { color:var(--teal); font-size:0.7rem; letter-spacing:0.1em; text-transform:uppercase; font-weight:600; margin:0 0 0.5rem; font-family:'Syne',sans-serif; }
.rcba-login-card p { color:var(--muted); font-size:0.84rem; margin:0 0 1.8rem; line-height:1.6; }
.rcba-gbtn { display:inline-flex; align-items:center; gap:0.6rem; background:rgba(255,255,255,0.06); color:var(--text); border:1px solid rgba(255,255,255,0.12); border-radius:12px; padding:11px 26px; font-size:0.9rem; font-weight:500; text-decoration:none; backdrop-filter:blur(10px); font-family:'Outfit',sans-serif; transition:all .2s; }
.rcba-gbtn:hover { background:rgba(255,255,255,0.1); border-color:rgba(255,255,255,0.2); transform:translateY(-1px); }
.rcba-denied { background:rgba(255,107,107,0.08); border:1px solid rgba(255,107,107,0.2); border-radius:14px; padding:2rem; text-align:center; }
.rcba-denied h3 { color:var(--danger); font-family:'Syne',sans-serif; }
.rcba-denied p  { color:var(--muted); font-size:0.88rem; }

/* STREAMLIT OVERRIDES */
.stButton > button { font-family:'Outfit',sans-serif !important; border-radius:12px !important; font-weight:500 !important; transition:all .2s !important; }
.stButton > button[kind="primary"] { background:linear-gradient(135deg,var(--teal),var(--teal-dark)) !important; border:none !important; color:var(--navy) !important; font-weight:700 !important; font-family:'Syne',sans-serif !important; box-shadow:0 0 20px var(--teal-glow) !important; }
.stButton > button[kind="secondary"] { background:transparent !important; border:1px solid var(--border) !important; color:var(--muted) !important; }
.stTextInput > div > div > input, .stTextArea > div > div > textarea { background:var(--surface2) !important; border:1px solid var(--border) !important; border-radius:10px !important; color:var(--text) !important; font-family:'Outfit',sans-serif !important; font-size:0.9rem !important; }
.stSelectbox > div > div { background:var(--surface2) !important; border:1px solid var(--border) !important; border-radius:10px !important; color:var(--text) !important; }
label { font-family:'Outfit',sans-serif !important; font-size:0.76rem !important; color:var(--muted) !important; font-weight:500 !important; letter-spacing:0.05em !important; text-transform:uppercase !important; }
.stAlert { border-radius:12px !important; }
section[data-testid="stSidebar"] { display:none !important; }
#MainMenu, footer { visibility:hidden; }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def word_count(text: str) -> int:
    return len(text.split()) if text.strip() else 0


def truncate_words(text: str, limit: int) -> str:
    words = text.split()
    return " ".join(words[:limit]) if len(words) > limit else text


def word_counter_html(text, limit):
    wc  = word_count(text)
    cls = "over" if wc > limit else ("near" if wc > limit * 0.85 else "")
    return f"<div class='word-counter {cls}'>{wc} / {limit} words</div>"


def load_members():
    app_dir = os.path.dirname(os.path.abspath(__file__))
    members_file = os.path.join(app_dir, "members.json")
    # Prefer Supabase
    try:
        from supabase_handler import get_all_members, SUPABASE_ENABLED
        if SUPABASE_ENABLED:
            rows = get_all_members()
            if rows:
                return rows
    except Exception:
        pass
    # JSON fallback
    if os.path.exists(members_file):
        try:
            with open(members_file) as f:
                return json.load(f)
        except Exception:
            pass
    return []


def render_header():
    logo_html = (
        f"<div class='rcba-logo-wrap'><img src='{LOGO_SRC}' alt='RCBA'>"
        f"<div class='rcba-logo-ring'></div><div class='rcba-logo-ring2'></div></div>"
        if LOGO_SRC else
        "<div class='rcba-logo-wrap' style='background:linear-gradient(135deg,#00C9B1,#00A896);"
        "border-radius:50%;display:flex;align-items:center;justify-content:center;"
        "font-size:1.8rem;color:#0A0F1E;font-family:Syne;font-weight:800;'>R</div>"
    )
    st.markdown(f"""
    <div class='rcba-header'>
        <div class='rcba-header-stripe'></div>
        <div class='rcba-header-body'>
            {logo_html}
            <div class='rcba-header-text'>
                <h1>Rotaract Club of Bombay Airport</h1>
                <p class='sub'>Event Management System</p>
                <p class='tagline'>Act For Impact</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_navbar(active_page: str, role: str, name: str, picture: str):
    init    = name[0].upper() if name else "M"
    av_html = (f"<img class='rcba-nav-avatar' src='{picture}' alt='avatar'>"
               if picture else f"<span class='rcba-nav-avatar-init'>{init}</span>")

    role_label = role_display_label(role)
    role_cls   = role_badge_class(role)

    st.markdown(f"""
    <div class='rcba-nav'>
        <div class='rcba-nav-links'></div>
        <div class='rcba-nav-right'>
            <span class='rcba-role-badge {role_cls}'>{role_label}</span>
            {av_html}
        </div>
    </div>
    """, unsafe_allow_html=True)

    can_manage = role in ("admin", "secretariat")
    is_admin_r = role == "admin"

    # Build column layout
    if can_manage and is_admin_r:
        cols = st.columns([1.2, 1.3, 1.1, 1.1, 0.95, 0.8])
        sign_out_col = 5
    elif can_manage:
        cols = st.columns([1.2, 1.3, 1.1, 1.1, 0.8])
        sign_out_col = 4
    else:
        cols = st.columns([1.2, 1.3, 1.1, 0.8])
        sign_out_col = 3

    with cols[0]:
        if st.button("Dashboard", key="nav_dashboard", use_container_width=True, type="secondary"):
            st.session_state.page = "dashboard"
            st.rerun()
    with cols[1]:
        if st.button("New Report", key="nav_new_report", use_container_width=True, type="secondary"):
            st.session_state.page = "new_report"
            st.rerun()
    with cols[2]:
        if st.button("Profile", key="nav_profile", use_container_width=True, type="secondary"):
            st.session_state.page = "profile"
            st.rerun()

    col_offset = 3
    if can_manage:
        with cols[col_offset]:
            if st.button("Members", key="nav_members", use_container_width=True, type="secondary"):
                st.session_state.page = "manage_members"
                st.rerun()
        col_offset += 1

    if is_admin_r:
        with cols[col_offset]:
            if st.button("Admin", key="nav_admin", use_container_width=True, type="secondary"):
                st.session_state.page = "admin"
                st.rerun()

    with cols[sign_out_col]:
        if st.button("Sign Out", key="nav_signout", use_container_width=True, type="secondary"):
            logout()


def render_step(num, title):
    st.markdown(f"<div class='rcba-step'><span class='rcba-step-pill'>Step {num}</span>"
                f"<p class='rcba-step-title'>{title}</p></div>", unsafe_allow_html=True)


def render_hr():
    st.markdown("<div class='rcba-hr'></div>", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# GOOGLE OAUTH
# ═══════════════════════════════════════════════════════════════════════════════

def build_auth_url(state: str) -> str:
    p = {
        "client_id":     GOOGLE_CLIENT_ID,
        "redirect_uri":  REDIRECT_URI,
        "response_type": "code",
        "scope":         "openid email profile",
        "state":         state,
        "access_type":   "online",
        "prompt":        "select_account",
    }
    return GOOGLE_AUTH_URL + "?" + urllib.parse.urlencode(p)


def exchange_code(code: str) -> dict:
    r = req.post(GOOGLE_TOKEN_URL, data={
        "code": code, "client_id": GOOGLE_CLIENT_ID,
        "client_secret": GOOGLE_CLIENT_SECRET,
        "redirect_uri": REDIRECT_URI, "grant_type": "authorization_code",
    }, timeout=15)
    r.raise_for_status()
    return r.json()


def fetch_user(token: str) -> dict:
    r = req.get(GOOGLE_INFO_URL, headers={"Authorization": f"Bearer {token}"}, timeout=15)
    r.raise_for_status()
    return r.json()


def is_whitelisted(email: str) -> bool:
    return email.strip().lower() in WHITELISTED_EMAILS


def logout():
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# LOGIN SCREEN
# ═══════════════════════════════════════════════════════════════════════════════

def login_screen():
    render_header()

    problems = []
    if not has_google_credentials():       problems.append("`GOOGLE_CREDENTIALS_JSON` secret or `google_credentials.json` file not found")
    if not GOOGLE_CLIENT_ID:               problems.append("`client_id` missing from credentials")
    if not GOOGLE_CLIENT_SECRET:           problems.append("`client_secret` missing from credentials")
    if not WHITELISTED_EMAILS:             problems.append("`WHITELISTED_EMAILS` not set")

    if problems:
        st.error("Configuration issues:\n\n" + "\n".join(f"- {p}" for p in problems))
        with st.expander("Setup instructions"):
            st.markdown(f"""
**Render Setup:**
1. Go to your Render service → Environment → Add environment variable:
   - `GOOGLE_CREDENTIALS_JSON` = paste your Google OAuth JSON (the whole file)
   - `WHITELISTED_EMAILS`      = comma-separated emails
   - `OAUTH_REDIRECT_URI`      = your Render app URL ending in `/` (e.g. `https://my-app.onrender.com/`)
   - `SUPABASE_URL` and `SUPABASE_KEY`
2. Redeploy.

**Local Setup:**
1. Create `google_credentials.json` beside `app.py`
2. Create `.env` with `WHITELISTED_EMAILS=...` and other secrets
3. `pip install -r requirements.txt && streamlit run app.py`
""")
        return

    # ── Process OAuth callback ─────────────────────────────────────────────
    auth_code = st.session_state.get("pending_code", "")
    if auth_code:
        with st.spinner("Signing you in…"):
            try:
                td = exchange_code(auth_code)
                at = td.get("access_token", "")
                if not at:
                    raise ValueError("No access token returned.")
                ui = fetch_user(at)
            except Exception as exc:
                st.session_state.pending_code = ""
                st.error(f"Sign-in failed: {exc}")
                return

        st.session_state.pending_code = ""
        email   = ui.get("email", "").lower()
        name    = ui.get("name", email)
        picture = ui.get("picture", "")

        if is_whitelisted(email):
            role = get_role(email)
            st.session_state.update(
                logged_in=True, username=name, user_email=email,
                user_picture=picture, role=role, page="dashboard"
            )
            st.rerun()
        else:
            _, col, _ = st.columns([1, 2, 1])
            with col:
                st.markdown(
                    f"<div class='rcba-denied'><h3>Access Denied</h3>"
                    f"<p><strong>{email}</strong> is not on the authorised member list.</p></div>",
                    unsafe_allow_html=True
                )
        return

    # ── Render login button ────────────────────────────────────────────────
    state    = secrets.token_hex(16)
    st.session_state["oauth_state"] = state        # store for CSRF check on return
    auth_url = build_auth_url(state)

    logo_html = f"<img src='{LOGO_SRC}' alt='RCBA'>" if LOGO_SRC else ""
    st.markdown(f"""
    <div class='rcba-login-outer'>
        <div class='rcba-login-card'>
            {logo_html}
            <p class='club-tag'>Rotaract Club of Bombay Airport</p>
            <h2>Event Reporter</h2>
            <p>Sign in with your Google account.<br>Access restricted to authorised RCBA members.</p>
            <a href='{auth_url}' target='_self' class='rcba-gbtn'>
                <img src='https://www.google.com/favicon.ico' width='16' alt='G'>
                Continue with Google
            </a>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# EMAIL — review notification
# ═══════════════════════════════════════════════════════════════════════════════

def send_review_email(report: dict, action: str, comment: str, reviewer_name: str) -> None:
    if not SENDER_EMAIL or not SENDER_PASSWORD:
        raise ValueError("SENDER_EMAIL and SENDER_PASSWORD must be set.")

    to_email    = report.get("submitted_by_email", "")
    to_name     = report.get("submitted_by_name", "Member")
    event_title = report.get("event_title", "Your Report")

    if not to_email:
        raise ValueError("Submitter email not found in report record.")

    clean_from = SENDER_EMAIL.strip().strip('"\'')
    clean_pass = SENDER_PASSWORD.strip().strip('"\'').replace(" ", "")

    action_word = "Approved" if action == "approve" else "Rejected"
    subject     = f"RCBA Report {action_word}: {event_title}"

    if action == "approve":
        body = (
            f"Dear {to_name},\n\n"
            f"Your event report for '{event_title}' has been APPROVED.\n\n"
            f"Reviewed by : {reviewer_name}\n"
        )
        if comment.strip():
            body += f"Comments    : {comment.strip()}\n"
        body += f"\nGreat work!\n\nRegards,\n{reviewer_name}\nRotaract Club of Bombay Airport"
    else:
        body = (
            f"Dear {to_name},\n\n"
            f"Your event report for '{event_title}' has been REJECTED.\n\n"
            f"Reviewed by      : {reviewer_name}\n"
            f"Rejection Reason :\n{comment.strip() or 'Please check with the reviewer.'}\n\n"
            f"Please make the necessary revisions and resubmit.\n\n"
            f"Regards,\n{reviewer_name}\nRotaract Club of Bombay Airport"
        )

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"]    = f"{reviewer_name} <{clean_from}>"
    msg["To"]      = to_email
    msg.set_content(body)

    sent = False
    last_err = None
    for method in ["starttls", "ssl"]:
        try:
            if method == "starttls":
                with smtplib.SMTP("smtp.gmail.com", 587, timeout=20) as srv:
                    srv.ehlo(); srv.starttls(); srv.ehlo()
                    srv.login(clean_from, clean_pass)
                    srv.send_message(msg)
            else:
                with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=20) as srv:
                    srv.login(clean_from, clean_pass)
                    srv.send_message(msg)
            sent = True
            break
        except smtplib.SMTPAuthenticationError:
            raise RuntimeError(
                "Gmail auth failed. Use a 16-char App Password (no spaces).\n"
                "Generate at: myaccount.google.com/apppasswords"
            )
        except Exception as exc:
            last_err = exc
    if not sent:
        raise RuntimeError(f"Could not connect to Gmail: {last_err}")


# ═══════════════════════════════════════════════════════════════════════════════
# EMAIL — secretary submission
# ═══════════════════════════════════════════════════════════════════════════════

def send_email(event_title: str, bod: dict, docx_bytes: bytes) -> None:
    if not all([SENDER_EMAIL, SENDER_PASSWORD, SECRETARY_EMAIL]):
        raise ValueError("SENDER_EMAIL, SENDER_PASSWORD, SECRETARY_EMAIL must all be set.")

    name       = st.session_state.get("username", "Member")
    user_email = st.session_state.get("user_email", "")
    clean_from = SENDER_EMAIL.strip().strip('"\'')
    clean_pass = SENDER_PASSWORD.strip().strip('"\'').replace(" ", "")
    clean_to   = SECRETARY_EMAIL.strip().strip('"\'')

    msg = EmailMessage()
    msg["Subject"] = f"RCBA Event Report \u2014 {event_title}"
    msg["From"]    = f"{name} <{clean_from}>"
    msg["To"]      = clean_to

    avenue_chairs = bod.get("avenue_chairs", [])
    chair_str     = ", ".join(avenue_chairs) if isinstance(avenue_chairs, list) else str(avenue_chairs)

    msg.set_content(
        f"Dear Secretary,\n\nPlease find attached the event report for '{event_title}'.\n\n"
        f"Attendance    : {bod.get('total_attendance', 'N/A')}\n"
        f"Drive Link    : {bod.get('drive_link', 'N/A')}\n"
        f"Income        : INR {bod.get('income', 0)}\n"
        f"Expenditure   : INR {bod.get('expenditure', 0)}\n"
        f"Profit / Loss : INR {bod.get('profit_loss', 0)}\n"
        f"Avenue Chairs : {chair_str or 'N/A'}\n"
        f"Submitted by  : {name} <{user_email}>\n\n"
        f"Regards,\n{name}\nRotaract Club of Bombay Airport"
    )
    msg.add_attachment(
        docx_bytes,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"RCBA_Report_{event_title.replace(' ', '_')}.docx",
    )

    sent = False
    last_err = None
    for method in ["starttls", "ssl"]:
        try:
            if method == "starttls":
                with smtplib.SMTP("smtp.gmail.com", 587, timeout=20) as srv:
                    srv.ehlo(); srv.starttls(); srv.ehlo()
                    srv.login(clean_from, clean_pass)
                    srv.send_message(msg)
            else:
                with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=20) as srv:
                    srv.login(clean_from, clean_pass)
                    srv.send_message(msg)
            sent = True
            break
        except smtplib.SMTPAuthenticationError:
            raise RuntimeError(
                "Gmail auth failed. Use a 16-char App Password (no spaces).\n"
                "Generate at: myaccount.google.com/apppasswords"
            )
        except Exception as exc:
            last_err = exc
    if not sent:
        raise RuntimeError(f"Could not connect to Gmail: {last_err}")


# ═══════════════════════════════════════════════════════════════════════════════
# AI REPORT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def generate_report(event: dict) -> str:
    try:
        from groq import Groq
    except ImportError:
        raise ImportError("groq package not installed. Run: pip install groq")
    if not GROQ_API_KEY:
        raise ValueError("GROQ_API_KEY is not set.")

    client = Groq(api_key=GROQ_API_KEY)

    desc_words    = word_count(event.get("description", ""))
    pre_words     = word_count(event.get("pre_event", ""))
    on_words      = word_count(event.get("on_day", ""))
    post_words    = word_count(event.get("post_event", ""))
    outcome_words = word_count(event.get("outcome", ""))

    aim_budget    = min(max(desc_words * 3, 120), 300)
    pre_budget    = min(max(pre_words * 3, 120), 350)
    on_budget     = min(max(on_words * 3, 120), 350)
    post_budget   = min(max(post_words * 3, 80), 250)
    impact_budget = min(max(outcome_words * 3, 120), 300)

    prompt = f"""
You are a professional secretary for Rotaract Club of Bombay Airport (RCBA).
Write a formal event report in formal English for club records and Rotary District submissions.

EVENT DETAILS:
- Title       : {event['title']}
- Start       : {event.get('start_dt','N/A')}
- End         : {event.get('end_dt','N/A')}
- Venue       : {event['venue']}
- Chief Guest : {event['chief_guest']}
- Description : {truncate_words(event.get('description',''), WORD_LIMITS['description'])}
- Outcome     : {truncate_words(event.get('outcome',''), WORD_LIMITS['outcome'])}

EXECUTION INPUTS:
- Pre-Event Work  : {truncate_words(event.get('pre_event','N/A'), WORD_LIMITS['pre_event'])}
- On-Day Work     : {truncate_words(event.get('on_day','N/A'), WORD_LIMITS['on_day'])}
- Post-Event Work : {truncate_words(event.get('post_event','N/A'), WORD_LIMITS['post_event'])}

Generate a report with EXACTLY these four sections. Use plain headings — no #, **, or markdown.

SECTIONS AND WORD BUDGETS:
1. Aim (~{aim_budget} words)
2. Execution — with exactly three plain subheadings on their own lines:
     Pre-Event Work (~{pre_budget} words)
     On-Day Work (~{on_budget} words)
     Post-Event Work (~{post_budget} words)
3. Impact Analysis (~{impact_budget} words)
4. Follow Up and Feedback (~150 words)

Strictly follow the word budgets. Start directly with "Aim".
""".strip()

    for model in ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"]:
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.65,
                max_tokens=2800,
            )
            return resp.choices[0].message.content.strip()
        except Exception:
            continue
    raise RuntimeError("Report generation service unavailable. Please try again.")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_docx(event: dict, report_text: str, bod: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    TEAL  = RGBColor(0x2E, 0x7D, 0x74)
    GOLD  = RGBColor(0xC9, 0x82, 0x3A)
    NAVY  = RGBColor(0x1A, 0x2B, 0x3C)
    GREY  = RGBColor(0x5A, 0x6E, 0x7A)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
        sec.left_margin = Inches(1.1); sec.right_margin = Inches(1.1)

    def shade_cell(cell, hx):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hx); tcPr.append(shd)

    def border_bottom(para, color, sz="6"):
        pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom"); bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), sz); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), color)
        pBdr.append(bot); pPr.append(pBdr)

    logo_path = _logo_path()

    # ── header table ──────────────────────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=2); hdr.autofit = False
    hdr.columns[0].width = Inches(1.1); hdr.columns[1].width = Inches(5.3)
    if logo_path:
        c0 = hdr.cell(0, 0)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c0.paragraphs[0].add_run().add_picture(logo_path, width=Inches(0.95))
    c1 = hdr.cell(0, 1); pc = c1.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rc = pc.add_run("ROTARACT CLUB OF BOMBAY AIRPORT")
    rc.bold = True; rc.font.size = Pt(15); rc.font.color.rgb = NAVY
    ps = c1.add_paragraph(f"Act For Impact  \u2022  RI Zone 2A  \u2022  RI Year {RI_YEAR}")
    ps.runs[0].font.size = Pt(9); ps.runs[0].font.color.rgb = GREY; ps.runs[0].italic = True

    sep1 = doc.add_paragraph(); sep1.paragraph_format.space_before = Pt(6); sep1.paragraph_format.space_after = Pt(8)
    border_bottom(sep1, "C9823A", "12")

    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("EVENT REPORT"); tr.bold = True; tr.font.size = Pt(10); tr.font.color.rgb = TEAL

    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_after = Pt(12)
    er = ep.add_run(event["title"]); er.bold = True; er.font.size = Pt(17); er.font.color.rgb = NAVY

    # ── meta table ────────────────────────────────────────────────────────
    profit_loss   = bod.get("profit_loss", bod.get("income", 0) - bod.get("expenditure", 0))
    avenue_chairs = bod.get("avenue_chairs", [])
    chair_str     = ", ".join(avenue_chairs) if isinstance(avenue_chairs, list) else str(avenue_chairs)
    member_list   = bod.get("member_attendance", [])
    member_str    = ", ".join(member_list) if isinstance(member_list, list) else str(member_list)

    meta_rows = [
        ("Project Name",           event["title"]),
        ("Event Start",            event.get("start_dt", "N/A")),
        ("Event End",              event.get("end_dt", "N/A")),
        ("Venue",                  event["venue"]),
        ("Chief Guest",            event["chief_guest"]),
        ("Avenue",                 bod.get("avenue", "N/A")),
        ("Project Level / Freq.",  bod.get("project_level", "N/A")),
        ("Project Hours",          bod.get("project_hours", "N/A")),
        ("Total Man Hours",        str(bod.get("man_hours", "N/A"))),
        ("Avenue Chair(s)",        chair_str or "N/A"),
        ("Drive / Images Link",    bod.get("drive_link", "N/A")),
        ("Total Attendance",       str(bod.get("total_attendance", "N/A"))),
        ("Member Attendance",      f"{bod.get('member_attendance_count', 0)} — {member_str}"),
        ("Guest Attendance",       f"{bod.get('guest_attendance_count', 0)} — {bod.get('guest_names', 'N/A')}"),
        ("District Attendance",    f"{bod.get('district_attendance_count', 0)} — {bod.get('district_names', 'N/A')}"),
        ("Ambassadorial Attend.",  f"{bod.get('ambassadorial_attendance_count', 0)} — {bod.get('ambassadorial_club_names', 'N/A')}"),
        ("Income",                 f"INR {bod.get('income', 0)}"),
        ("Expenditure",            f"INR {bod.get('expenditure', 0)}"),
        ("Sponsorship",            f"INR {bod.get('sponsorship', 0)}"),
        ("Profit / Loss",          f"INR {profit_loss}"),
        ("Club",                   "Rotaract Club of Bombay Airport"),
        ("Zone / RI Year",         f"Zone 2A  |  RI Year {RI_YEAR}"),
        ("Submitted By",           st.session_state.get("username", "Member")),
        ("Email",                  st.session_state.get("user_email", "")),
        ("Report Date",            date.today().strftime("%B %d, %Y")),
    ]

    mt = doc.add_table(rows=len(meta_rows), cols=2); mt.style = "Table Grid"
    for i, row in enumerate(mt.rows):
        row.cells[0].width = Inches(2.1); row.cells[1].width = Inches(4.3)
        label, value = meta_rows[i]
        shade_cell(row.cells[0], "1A2B3C")
        shade_cell(row.cells[1], "F2F6F5" if i % 2 == 0 else "FFFFFF")
        lp = row.cells[0].paragraphs[0]
        lp.paragraph_format.space_before = Pt(3); lp.paragraph_format.space_after = Pt(3)
        lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(9.5); lr.font.color.rgb = WHITE
        vp = row.cells[1].paragraphs[0]
        vp.paragraph_format.space_before = Pt(3); vp.paragraph_format.space_after = Pt(3)
        vr = vp.add_run(str(value)); vr.font.size = Pt(9.5); vr.font.color.rgb = NAVY

    doc.add_paragraph()

    # ── report body ───────────────────────────────────────────────────────
    blocks: dict = {}; current = None; buf: list = []
    for line in report_text.splitlines():
        s = line.strip()
        matched = next((sec for sec in REPORT_SECTIONS
                        if s.lower() == sec.lower() or s.lower().startswith(sec.lower() + ":")), None)
        if matched:
            if current and buf: blocks[current] = "\n".join(buf).strip()
            current, buf = matched, []
        elif current is not None:
            buf.append(line)
    if current and buf: blocks[current] = "\n".join(buf).strip()

    for sec_title in REPORT_SECTIONS:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(14); hp.paragraph_format.space_after = Pt(4)
        hr2 = hp.add_run(sec_title.upper()); hr2.bold = True; hr2.font.size = Pt(11.5); hr2.font.color.rgb = TEAL
        border_bottom(hp, "2E7D74", "4")
        content = blocks.get(sec_title, "")

        if sec_title == "Execution" and content:
            sb: dict = {}; sc = None; sb_buf: list = []
            for line in content.splitlines():
                s = line.strip()
                ms = next((sub for sub in EXECUTION_SUBSECTIONS
                           if s.lower() == sub.lower() or s.lower().startswith(sub.lower() + ":")), None)
                if ms:
                    if sc and sb_buf: sb[sc] = "\n".join(sb_buf).strip()
                    sc, sb_buf = ms, []
                elif sc is not None:
                    sb_buf.append(line)
            if sc and sb_buf: sb[sc] = "\n".join(sb_buf).strip()

            for sub_title in EXECUTION_SUBSECTIONS:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(10); sp.paragraph_format.space_after = Pt(3)
                sp.paragraph_format.left_indent = Inches(0.15)
                sr = sp.add_run(sub_title); sr.bold = True; sr.font.size = Pt(10.5); sr.font.color.rgb = GOLD
                for pt in sb.get(sub_title, "").split("\n\n"):
                    pt = pt.strip()
                    if pt:
                        p = doc.add_paragraph(pt); p.paragraph_format.space_after = Pt(5)
                        p.paragraph_format.left_indent = Inches(0.15)
                        for run in p.runs: run.font.size = Pt(10.5)
        elif content:
            for pt in content.split("\n\n"):
                pt = pt.strip()
                if pt:
                    p = doc.add_paragraph(pt); p.paragraph_format.space_after = Pt(5)
                    for run in p.runs: run.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph("[Content not generated]")
            p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA); p.runs[0].font.size = Pt(10.5)
        doc.add_paragraph()

    if bod.get("feedback", "").strip():
        hp2 = doc.add_paragraph()
        hr3 = hp2.add_run("MEMBER FEEDBACK"); hr3.bold = True; hr3.font.size = Pt(11.5); hr3.font.color.rgb = TEAL
        border_bottom(hp2, "2E7D74", "4")
        fbp = doc.add_paragraph(bod["feedback"])
        fbp.runs[0].font.size = Pt(10.5); fbp.runs[0].italic = True; fbp.runs[0].font.color.rgb = GREY
        doc.add_paragraph()

    sep2 = doc.add_paragraph(); sep2.paragraph_format.space_before = Pt(6)
    border_bottom(sep2, "C9823A", "8")
    sig = doc.add_paragraph(); sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_r = sig.add_run(
        f"Prepared by: {st.session_state.get('username','Member')}\n"
        f"Email: {st.session_state.get('user_email','')}\n"
        f"Date: {date.today().strftime('%B %d, %Y')}"
    )
    sig_r.font.size = Pt(9); sig_r.font.color.rgb = GREY; sig_r.italic = True

    out = io.BytesIO(); doc.save(out); out.seek(0)
    return out.read()


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════

def page_dashboard():
    role = st.session_state.get("role", "director")
    if role in ("secretariat", "admin"):
        page_dashboard_secretariat(send_review_email)
    elif role == "editor":
        page_dashboard_editor(send_review_email)
    else:
        page_dashboard_director()


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: PROFILE
# ═══════════════════════════════════════════════════════════════════════════════

def page_profile():
    name    = st.session_state.get("username", "Member")
    email   = st.session_state.get("user_email", "")
    picture = st.session_state.get("user_picture", "")
    role    = st.session_state.get("role", "director")

    init       = name[0].upper() if name else "M"
    role_label = role_display_label(role)
    role_cls   = role_badge_class(role)

    av_html = (
        f"<img class='rcba-profile-avatar-lg' src='{picture}' alt='avatar'>"
        if picture else
        f"<div class='rcba-profile-avatar-lg-init'>{init}</div>"
    )

    all_reports = load_reports(email=email, role=role)
    my_reports  = [r for r in all_reports if r.get("submitted_by_email", "").lower() == email.lower()]
    total       = len(my_reports)
    late_count  = sum(1 for r in my_reports if is_late_report(r))
    approved_ct = sum(1 for r in my_reports if get_status(r) == "Approved")
    pending_ct  = sum(1 for r in my_reports if get_status(r) == "Pending")

    st.markdown(f"""
    <div class='rcba-profile-hero'>
        {av_html}
        <p class='rcba-profile-name'>{name}</p>
        <p class='rcba-profile-email'>{email}</p>
        <span class='rcba-role-badge {role_cls}' style='margin-bottom:1rem;display:inline-block;'>{role_label}</span>
        <div class='rcba-profile-stats'>
            <div class='rcba-profile-stat'><span class='n'>{total}</span><span class='l'>Submitted</span></div>
            <div class='divider'></div>
            <div class='rcba-profile-stat'><span class='n' style='color:var(--success)'>{approved_ct}</span><span class='l'>Approved</span></div>
            <div class='divider'></div>
            <div class='rcba-profile-stat'><span class='n' style='color:var(--gold)'>{pending_ct}</span><span class='l'>Pending</span></div>
            <div class='divider'></div>
            <div class='rcba-profile-stat'><span class='n' style='color:var(--danger)'>{late_count}</span><span class='l'>Late</span></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<div class='rcba-card'><div class='rcba-section-label' style='margin-bottom:0.6rem;'>Account Details</div>", unsafe_allow_html=True)
    for key, val in [
        ("Full Name",     name),
        ("Email Address", email),
        ("Role",          role_label),
        ("Club",          "Rotaract Club of Bombay Airport"),
        ("RI Year",       RI_YEAR),
    ]:
        st.markdown(f"<div class='rcba-info-row'><span class='key'>{key}</span><span class='val'>{val}</span></div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: NEW REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def page_new_report():
    role = st.session_state.get("role", "director")
    if not can_submit_report(role):
        st.error("You don't have permission to submit reports.")
        return

    st.markdown("<p class='rcba-page-title'>New Event Report</p>", unsafe_allow_html=True)
    st.markdown("<p class='rcba-page-sub'>Fill in the details below to generate and submit your report.</p>", unsafe_allow_html=True)

    render_step("01", "Event Details")
    st.markdown("<div class='rcba-card'>", unsafe_allow_html=True)

    # Basic event fields (outside form so word counters update live on each keystroke)
    c1, c2 = st.columns(2)
    with c1:
        title       = st.text_input("Event Title *",  placeholder="e.g. Shukriya Award Night 2025", key="f_title")
        venue       = st.text_input("Venue *",         placeholder="e.g. Sanskriti Arts, Andheri", key="f_venue")
    with c2:
        chief_guest = st.text_input("Chief Guest",     placeholder="Name or N/A", key="f_guest")

    st.markdown("<div class='rcba-section-label' style='margin-top:0.8rem;'>Event Schedule</div>", unsafe_allow_html=True)
    dc1, dc2 = st.columns(2)
    with dc1:
        start_date = st.date_input("Start Date *",  value=date.today(), key="f_sdate")
        start_time = st.time_input("Start Time *",  value=datetime.strptime("18:00", "%H:%M").time(), key="f_stime")
    with dc2:
        end_date   = st.date_input("End Date *",    value=date.today(), key="f_edate")
        end_time   = st.time_input("End Time *",    value=datetime.strptime("20:00", "%H:%M").time(), key="f_etime")

    st.markdown("<div class='rcba-section-label' style='margin-top:1rem;'>Description</div>", unsafe_allow_html=True)
    description = st.text_area(f"Event Description / Aim * (max {WORD_LIMITS['description']} words)", placeholder="Purpose and motivation behind this event?", height=85, key="f_desc")
    st.markdown(word_counter_html(description, WORD_LIMITS["description"]), unsafe_allow_html=True)

    st.markdown("<div class='rcba-section-label' style='margin-top:1rem;'>Execution Details</div>", unsafe_allow_html=True)

    pre_event = st.text_area(f"Pre-Event Work * (max {WORD_LIMITS['pre_event']} words)", placeholder="Planning, logistics, venue booking, coordination...", height=80, key="f_pre")
    st.markdown(word_counter_html(pre_event, WORD_LIMITS["pre_event"]), unsafe_allow_html=True)

    on_day = st.text_area(f"On-Day Work * (max {WORD_LIMITS['on_day']} words)", placeholder="Schedule, segments, hosting, key moments...", height=80, key="f_on")
    st.markdown(word_counter_html(on_day, WORD_LIMITS["on_day"]), unsafe_allow_html=True)

    post_event = st.text_area(f"Post-Event Work (max {WORD_LIMITS['post_event']} words)", placeholder="Follow-up tasks, thank-you messages, social media... (optional)", height=70, key="f_post")
    st.markdown(word_counter_html(post_event, WORD_LIMITS["post_event"]), unsafe_allow_html=True)

    st.markdown("<div class='rcba-section-label' style='margin-top:1rem;'>Outcome</div>", unsafe_allow_html=True)
    outcome = st.text_area(f"Outcome / Impact * (max {WORD_LIMITS['outcome']} words)", placeholder="Results, achievements, effect on members and community...", height=80, key="f_out")
    st.markdown(word_counter_html(outcome, WORD_LIMITS["outcome"]), unsafe_allow_html=True)

    gen_btn = st.button("Generate Report", use_container_width=True, type="primary", key="btn_gen_report")

    st.markdown("</div>", unsafe_allow_html=True)

    if gen_btn:
        start_dt = datetime.combine(start_date, start_time)
        end_dt   = datetime.combine(end_date, end_time)
        required = {"Event Title": title, "Venue": venue, "Description": description,
                    "Pre-Event Work": pre_event, "On-Day Work": on_day, "Outcome": outcome}
        missing = [k for k, v in required.items() if not v.strip()]
        if missing:
            st.error(f"Please fill in: {', '.join(missing)}")
        elif end_dt <= start_dt:
            st.error("End date/time must be after start date/time.")
        else:
            ev = {
                "title":       title.strip(),
                "start_dt":    start_dt.strftime("%d %b %Y, %I:%M %p"),
                "end_dt":      end_dt.strftime("%d %b %Y, %I:%M %p"),
                "venue":       venue.strip(),
                "chief_guest": chief_guest.strip() or "N/A",
                "description": description.strip(),
                "pre_event":   pre_event.strip(),
                "on_day":      on_day.strip(),
                "post_event":  post_event.strip(),
                "outcome":     outcome.strip(),
                "start_date":  start_date,
                "end_date":    end_date,
            }
            st.session_state.event       = ev
            st.session_state.report_text = ""
            st.session_state.docx_bytes  = None

            with st.spinner("Writing your report with AI…"):
                try:
                    st.session_state.report_text = generate_report(ev)
                    st.success("Report ready — review below, then fill in project details.")
                except Exception as exc:
                    st.error(f"Report generation failed: {exc}")

    if st.session_state.get("report_text"):
        render_hr()
        render_step("02", "Review & Edit Report")

        st.markdown("""
        <div class='rcba-edit-banner'>
            <h4>Your report is ready</h4>
            <p>Review and make any edits before building the document.</p>
        </div>
        """, unsafe_allow_html=True)

        edited = st.text_area(
            label="_re",
            value=st.session_state.report_text,
            height=540,
            label_visibility="collapsed",
            key="report_edit_area",
        )
        st.session_state.report_text = edited

        render_hr()
        render_step("03", "Project Details")

        st.markdown("""
        <div class='rcba-bod'>
            <p class='rcba-bod-title'>Post-Event Information</p>
            <p class='rcba-bod-sub'>Fill these in after the event. All fields appear in the document.</p>
        </div>
        """, unsafe_allow_html=True)

        members      = load_members()
        member_names = [m.get("name", "") for m in members if m.get("name")]

        # ── Attendance ───────────────────────────────────────────────────
        st.markdown("<div class='rcba-section-label'>Attendance Details</div>", unsafe_allow_html=True)
        st.info("**Member Names** is mandatory. Other attendance fields are optional.")

        att1, att2 = st.columns([2, 1])
        with att1:
            selected_members = st.multiselect(
                "Select Member Names *", options=member_names, key="member_attendance_final",
                help="Select members who attended the event",
            )
        with att2:
            st.metric("Member Count", len(selected_members))

        att3, att4 = st.columns(2)
        with att3:
            guest_count = st.number_input("Guest Attendance Count", value=0, min_value=0, key="guest_count_final")
            guest_names = st.text_area("Guest Names (comma-separated)", key="guest_names_final", height=60) if guest_count > 0 else ""
        with att4:
            district_count = st.number_input("District Attendance Count", value=0, min_value=0, key="district_count_final")
            district_names = st.text_area("District Member Names (comma-separated)", key="district_names_final", height=60) if district_count > 0 else ""

        att5, att6 = st.columns(2)
        with att5:
            ambassador_count = st.number_input("Ambassadorial Attendance Count", value=0, min_value=0, key="ambassador_count_final")
            ambassador_clubs = st.text_area("Club Names (comma-separated)", key="club_names_final", height=60) if ambassador_count > 0 else ""
        with att6:
            total_attendance = len(selected_members) + guest_count + district_count + ambassador_count
            st.metric("Total Attendance", total_attendance)

        # ── Financial ────────────────────────────────────────────────────
        st.markdown("<div class='rcba-section-label' style='margin-top:1.5rem;'>Financial & Project Details</div>", unsafe_allow_html=True)
        bc1, bc2 = st.columns(2)
        with bc1:
            drive_link    = st.text_input("Drive / Images Link", placeholder="https://drive.google.com/…", key="b_drive")
            avenue_chairs = st.multiselect("Avenue Chair(s) *", options=member_names, key="b_ave_chair")
            income        = st.number_input("Income (INR)",      min_value=0, value=0, step=100, key="b_inc")
            sponsorship   = st.number_input("Sponsorship (INR)", min_value=0, value=0, step=100, key="b_spon")
            avenue        = st.text_input("Avenue", placeholder="e.g. Club Service", key="b_ave")
        with bc2:
            expenditure   = st.number_input("Expenditure (INR)",       min_value=0, value=0, step=100, key="b_exp")
            project_level = st.text_input("Project Level / Frequency", placeholder="e.g. Club Project / Once", key="b_plvl")
            project_hours = st.text_input("Project Hours",             placeholder="e.g. 2", key="b_ph")

        profit_loss = income - expenditure
        pl_class    = "rcba-auto-negative" if profit_loss < 0 else ""

        ev_data    = st.session_state.get("event", {})
        s_date     = ev_data.get("start_date", date.today())
        e_date     = ev_data.get("end_date",   date.today())
        num_days   = max((e_date - s_date).days + 1, 1)
        man_hours  = num_days * 24

        ac1, ac2 = st.columns(2)
        with ac1:
            st.markdown(f"<div class='rcba-auto-field {pl_class}'><span class='label'>Profit / Loss (auto)</span><span class='val'>INR {profit_loss:,}</span></div>", unsafe_allow_html=True)
        with ac2:
            st.markdown(f"<div class='rcba-auto-field'><span class='label'>Total Man Hours (days × 24)</span><span class='val'>{man_hours}</span></div>", unsafe_allow_html=True)

        feedback = st.text_area("Member Feedback / Quotes (optional)", placeholder="Notable quotes from members…", height=70, key="b_fb")

        bod = {
            "drive_link":                    drive_link,
            "total_attendance":              total_attendance,
            "member_attendance":             selected_members,
            "member_attendance_count":       len(selected_members),
            "guest_attendance_count":        guest_count,
            "guest_names":                   guest_names,
            "district_attendance_count":     district_count,
            "district_names":                district_names,
            "ambassadorial_attendance_count": ambassador_count,
            "ambassadorial_club_names":      ambassador_clubs,
            "avenue_chairs":                 avenue_chairs,
            "income":                        income,
            "expenditure":                   expenditure,
            "sponsorship":                   sponsorship,
            "profit_loss":                   profit_loss,
            "avenue":                        avenue,
            "project_hours":                 project_hours,
            "man_hours":                     man_hours,
            "project_level":                 project_level,
            "feedback":                      feedback,
        }
        st.session_state.bod = bod

        render_hr()
        render_step("04", "Generate Document")

        if st.button("Build Document", use_container_width=True, type="primary"):
            with st.spinner("Building your document…"):
                try:
                    docx_bytes = build_docx(
                        st.session_state.event,
                        st.session_state.report_text,
                        st.session_state.bod,
                    )
                    st.session_state.docx_bytes = docx_bytes

                    ev_data  = st.session_state.event
                    ev_title = ev_data.get("title", "")
                    event_start_date_str = str(ev_data.get("start_date", date.today()))

                    record = {
                        "event_title":                   ev_title,
                        "avenue":                        bod.get("avenue", ""),
                        "avenue_chairs":                 bod.get("avenue_chairs", []),
                        "drive_link":                    bod.get("drive_link", ""),
                        "event_start_date":              event_start_date_str,
                        "submitted_by_name":             st.session_state.get("username", ""),
                        "submitted_by_email":            st.session_state.get("user_email", ""),
                        "submitted_at":                  str(datetime.now()),
                        "submission_timestamp":          str(datetime.now()),
                        "status":                        "submitted",
                        "approval_status":               "pending",
                        "role":                          st.session_state.get("role", "director"),
                        "rejection_message":             "",
                        "is_late":                       is_late(event_start_date_str, str(date.today())),
                        # attendance (canonical keys)
                        "total_attendance":              bod.get("total_attendance", 0),
                        "member_attendance":             bod.get("member_attendance", []),
                        "member_attendance_count":       bod.get("member_attendance_count", 0),
                        "guest_attendance_count":        bod.get("guest_attendance_count", 0),
                        "guest_names":                   bod.get("guest_names", ""),
                        "district_attendance_count":     bod.get("district_attendance_count", 0),
                        "district_names":                bod.get("district_names", ""),
                        "ambassadorial_attendance_count": bod.get("ambassadorial_attendance_count", 0),
                        "ambassadorial_club_names":      bod.get("ambassadorial_club_names", ""),
                    }
                    save_result = save_report(record, docx_binary=docx_bytes)
                    if save_result.get("success"):
                        st.success("Document ready to download. Report saved to dashboard.")
                    else:
                        st.warning(f"Report saved locally (Supabase error: {save_result.get('error', 'unknown')}). Contact admin if this persists.")
                except Exception as exc:
                    st.error(f"Document generation failed: {exc}")

        if st.session_state.get("docx_bytes"):
            ev_title = st.session_state.event["title"]
            st.download_button(
                label="Download Report (.docx)",
                data=st.session_state.docx_bytes,
                file_name=f"RCBA_Report_{ev_title.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, type="primary",
            )

            render_hr()
            render_step("05", "Send to Secretary")

            st.markdown(
                f"<p style='font-size:0.82rem;color:var(--muted);margin-bottom:1rem;line-height:1.6;'>"
                f"Will be sent to <strong style='color:var(--text);'>{SECRETARY_EMAIL or '[SECRETARY_EMAIL not set]'}</strong>"
                f" from <strong style='color:var(--text);'>{SENDER_EMAIL or '[SENDER_EMAIL not set]'}</strong>.</p>",
                unsafe_allow_html=True,
            )

            if st.button("Send Report by Email", use_container_width=True, type="primary"):
                with st.spinner("Sending…"):
                    try:
                        send_email(
                            event_title=ev_title,
                            bod=st.session_state.bod,
                            docx_bytes=st.session_state.docx_bytes,
                        )
                        st.balloons()
                        st.success(f"Sent to {SECRETARY_EMAIL}.")
                    except RuntimeError as exc:
                        st.error(str(exc))
                    except Exception as exc:
                        st.error(f"Email failed: {exc}")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN ROUTER
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    defaults = {
        "logged_in": False, "report_text": "", "docx_bytes": None,
        "event": {}, "bod": {}, "user_email": "", "user_picture": "",
        "username": "", "role": "director", "pending_code": "",
        "page": "dashboard", "oauth_state": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ── Handle OAuth callback ─────────────────────────────────────────────
    url_code  = st.query_params.get("code", "")
    url_state = st.query_params.get("state", "")

    if url_code and not st.session_state.pending_code:
        # CSRF check: state must match what we stored
        expected_state = st.session_state.get("oauth_state", "")
        if expected_state and url_state and url_state != expected_state:
            st.error("OAuth state mismatch — possible CSRF. Please try signing in again.")
            st.session_state.oauth_state = ""
            st.query_params.clear()
            st.rerun()
        else:
            st.session_state.pending_code = url_code
            st.session_state.oauth_state  = ""
            st.query_params.clear()
            st.rerun()

    if not st.session_state.logged_in:
        login_screen()
        return

    name    = st.session_state.get("username", "Member")
    email   = st.session_state.get("user_email", "")
    picture = st.session_state.get("user_picture", "")
    role    = st.session_state.get("role", "director")
    page    = st.session_state.get("page", "dashboard")

    render_header()
    render_navbar(active_page=page, role=role, name=name, picture=picture)

    if page == "dashboard":
        page_dashboard()
    elif page == "new_report":
        page_new_report()
    elif page == "profile":
        page_profile()
    elif page == "manage_members":
        if role in ("admin", "secretariat"):
            page_manage_members()
        else:
            st.error("Access denied.")
    elif page == "admin":
        if role == "admin":
            page_admin()
        else:
            st.error("Access denied.")


if __name__ == "__main__":
    main()
